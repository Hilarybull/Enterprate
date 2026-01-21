"""Product/Service Catalogue Service"""
import uuid
import io
import csv
import re
import os
import json
from datetime import datetime, timezone
from typing import List, Optional, Dict, Any
from fastapi import HTTPException, UploadFile
from pydantic import BaseModel
from app.core.database import get_db

# Intelligence Graph logging
from app.services.intelligence_service import log_catalogue_event

# Try to import pandas for Excel support
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

# Try to import Gemini for AI extraction
try:
    import google.generativeai as genai
    GEMINI_KEY = os.environ.get("GEMINI_API_KEY")
    GEMINI_AVAILABLE = bool(GEMINI_KEY)
except ImportError:
    GEMINI_AVAILABLE = False
    GEMINI_KEY = None

# Try to import PDF reader
try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Try to import DOCX reader
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class CatalogueItemCreate(BaseModel):
    name: str
    description: Optional[str] = None
    unitPrice: float
    currency: str = "GBP"
    taxRate: Optional[float] = None
    sku: Optional[str] = None
    category: Optional[str] = None


class CatalogueItemUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    unitPrice: Optional[float] = None
    currency: Optional[str] = None
    taxRate: Optional[float] = None
    sku: Optional[str] = None
    category: Optional[str] = None


class CatalogueService:
    """Service for managing product/service catalogue"""
    
    @staticmethod
    async def create_item(workspace_id: str, user_id: str, data: CatalogueItemCreate) -> dict:
        """Create a new catalogue item"""
        db = get_db()
        
        item_id = str(uuid.uuid4())
        now = datetime.now(timezone.utc).isoformat()
        
        item = {
            "id": item_id,
            "workspace_id": workspace_id,
            "name": data.name,
            "description": data.description,
            "unitPrice": data.unitPrice,
            "currency": data.currency,
            "taxRate": data.taxRate,
            "sku": data.sku,
            "category": data.category,
            "created_by": user_id,
            "createdAt": now,
            "updatedAt": now
        }
        
        await db.catalogue_items.insert_one(item)
        
        # Log event to intelligence graph
        await db.intelligence_events.insert_one({
            "id": str(uuid.uuid4()),
            "workspace_id": workspace_id,
            "type": "catalogue_item_added",
            "data": {
                "item_id": item_id,
                "name": data.name,
                "unitPrice": data.unitPrice,
                "category": data.category
            },
            "occurredAt": now
        })
        
        return {k: v for k, v in item.items() if k != '_id'}
    
    @staticmethod
    async def get_items(workspace_id: str, category: Optional[str] = None, search: Optional[str] = None) -> List[dict]:
        """Get all catalogue items for a workspace"""
        db = get_db()
        
        query = {"workspace_id": workspace_id}
        if category:
            query["category"] = category
        
        items = await db.catalogue_items.find(query).sort("name", 1).to_list(length=1000)
        
        result = [{k: v for k, v in item.items() if k != '_id'} for item in items]
        
        # Apply search filter if provided
        if search:
            search_lower = search.lower()
            result = [
                item for item in result
                if search_lower in (item.get("name") or "").lower()
                or search_lower in (item.get("sku") or "").lower()
                or search_lower in (item.get("category") or "").lower()
            ]
        
        return result
    
    @staticmethod
    async def get_item(item_id: str, workspace_id: str) -> dict:
        """Get a single catalogue item"""
        db = get_db()
        
        item = await db.catalogue_items.find_one({
            "id": item_id,
            "workspace_id": workspace_id
        })
        
        if not item:
            raise HTTPException(status_code=404, detail="Item not found")
        
        return {k: v for k, v in item.items() if k != '_id'}
    
    @staticmethod
    async def update_item(item_id: str, workspace_id: str, data: CatalogueItemUpdate) -> dict:
        """Update a catalogue item"""
        db = get_db()
        
        item = await db.catalogue_items.find_one({
            "id": item_id,
            "workspace_id": workspace_id
        })
        
        if not item:
            raise HTTPException(status_code=404, detail="Item not found")
        
        update_data = data.model_dump(exclude_unset=True)
        update_data["updatedAt"] = datetime.now(timezone.utc).isoformat()
        
        await db.catalogue_items.update_one(
            {"id": item_id, "workspace_id": workspace_id},
            {"$set": update_data}
        )
        
        updated = await db.catalogue_items.find_one({"id": item_id})
        return {k: v for k, v in updated.items() if k != '_id'}
    
    @staticmethod
    async def delete_item(item_id: str, workspace_id: str) -> dict:
        """Delete a catalogue item"""
        db = get_db()
        
        result = await db.catalogue_items.delete_one({
            "id": item_id,
            "workspace_id": workspace_id
        })
        
        if result.deleted_count == 0:
            raise HTTPException(status_code=404, detail="Item not found")
        
        return {"success": True, "message": "Item deleted"}
    
    @staticmethod
    async def process_upload(workspace_id: str, user_id: str, file: UploadFile) -> dict:
        """Process an uploaded catalogue file"""
        filename = file.filename.lower()
        content = await file.read()
        
        items = []
        
        try:
            if filename.endswith('.csv'):
                items = CatalogueService._parse_csv(content)
            elif filename.endswith(('.xlsx', '.xls')):
                items = CatalogueService._parse_excel(content)
            elif filename.endswith(('.pdf', '.doc', '.docx')):
                items = await CatalogueService._extract_from_document(content, filename)
            else:
                raise HTTPException(status_code=400, detail="Unsupported file format")
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse file: {str(e)}")
        
        # Validate items
        validation_results = CatalogueService._validate_items(items, workspace_id)
        
        # Log upload event
        db = get_db()
        await db.intelligence_events.insert_one({
            "id": str(uuid.uuid4()),
            "workspace_id": workspace_id,
            "type": "catalogue_uploaded",
            "data": {
                "filename": file.filename,
                "totalItems": len(items),
                "accepted": len(validation_results["accepted"]),
                "needsReview": len(validation_results["needsReview"]),
                "rejected": len(validation_results["rejected"])
            },
            "occurredAt": datetime.now(timezone.utc).isoformat()
        })
        
        return validation_results
    
    @staticmethod
    def _parse_csv(content: bytes) -> List[dict]:
        """Parse CSV content"""
        text = content.decode('utf-8-sig')  # Handle BOM
        reader = csv.DictReader(io.StringIO(text))
        
        items = []
        for row in reader:
            # Normalize column names
            normalized = {}
            for key, value in row.items():
                if key:
                    key_lower = key.lower().strip()
                    if 'name' in key_lower or key_lower == 'product' or key_lower == 'service':
                        normalized['name'] = value
                    elif 'desc' in key_lower:
                        normalized['description'] = value
                    elif 'price' in key_lower or key_lower == 'amount' or key_lower == 'cost':
                        normalized['unitPrice'] = value
                    elif 'currency' in key_lower:
                        normalized['currency'] = value
                    elif 'tax' in key_lower or 'vat' in key_lower:
                        normalized['taxRate'] = value
                    elif 'sku' in key_lower or 'code' in key_lower or 'id' in key_lower:
                        normalized['sku'] = value
                    elif 'category' in key_lower or 'type' in key_lower:
                        normalized['category'] = value
            
            if normalized.get('name'):
                items.append(normalized)
        
        return items
    
    @staticmethod
    def _parse_excel(content: bytes) -> List[dict]:
        """Parse Excel content"""
        if not PANDAS_AVAILABLE:
            raise HTTPException(status_code=400, detail="Excel parsing not available")
        
        df = pd.read_excel(io.BytesIO(content))
        
        # Normalize column names
        column_mapping = {}
        for col in df.columns:
            col_lower = str(col).lower().strip()
            if 'name' in col_lower or col_lower == 'product' or col_lower == 'service':
                column_mapping[col] = 'name'
            elif 'desc' in col_lower:
                column_mapping[col] = 'description'
            elif 'price' in col_lower or col_lower == 'amount':
                column_mapping[col] = 'unitPrice'
            elif 'currency' in col_lower:
                column_mapping[col] = 'currency'
            elif 'tax' in col_lower or 'vat' in col_lower:
                column_mapping[col] = 'taxRate'
            elif 'sku' in col_lower or 'code' in col_lower:
                column_mapping[col] = 'sku'
            elif 'category' in col_lower:
                column_mapping[col] = 'category'
        
        df = df.rename(columns=column_mapping)
        items = df.to_dict('records')
        
        # Clean up NaN values
        for item in items:
            for key, value in list(item.items()):
                if pd.isna(value):
                    item[key] = None
        
        return [item for item in items if item.get('name')]
    
    @staticmethod
    async def _extract_from_document(content: bytes, filename: str) -> List[dict]:
        """Extract catalogue items from PDF/Word using Gemini AI"""
        # First, extract text from the document
        text_content = ""
        
        if filename.endswith('.pdf'):
            if not PDF_AVAILABLE:
                raise HTTPException(
                    status_code=400,
                    detail="PDF parsing library not available. Please use CSV or Excel format."
                )
            try:
                pdf_reader = PdfReader(io.BytesIO(content))
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Failed to read PDF: {str(e)}")
        
        elif filename.endswith('.docx'):
            if not DOCX_AVAILABLE:
                raise HTTPException(
                    status_code=400,
                    detail="DOCX parsing library not available. Please use CSV or Excel format."
                )
            try:
                doc = DocxDocument(io.BytesIO(content))
                for para in doc.paragraphs:
                    text_content += para.text + "\n"
                # Also extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text for cell in row.cells)
                        text_content += row_text + "\n"
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Failed to read DOCX: {str(e)}")
        
        elif filename.endswith('.doc'):
            raise HTTPException(
                status_code=400,
                detail="Legacy .doc format not supported. Please convert to .docx or use CSV/Excel."
            )
        
        if not text_content.strip():
            raise HTTPException(status_code=400, detail="No text content found in document")
        
        # Use Gemini to extract structured data
        if not GEMINI_AVAILABLE:
            raise HTTPException(
                status_code=400,
                detail="AI extraction requires Gemini API key. Please use CSV or Excel format."
            )
        
        try:
            genai.configure(api_key=GEMINI_KEY)
            model = genai.GenerativeModel('gemini-2.0-flash')
            
            prompt = f"""Analyze the following text from a product/service catalogue and extract all items.
For each item, extract:
- name (required): The product or service name
- description: A brief description if available
- unitPrice (required): The price as a number (remove currency symbols)
- currency: The currency code (GBP, USD, EUR) - default to GBP if not specified
- taxRate: VAT/tax percentage if mentioned
- sku: Product code/SKU if available
- category: Product category if available

Return ONLY a valid JSON array of objects. Do not include any text before or after the JSON.
If no valid items can be extracted, return an empty array [].

Document text:
{text_content[:8000]}"""
            
            response = model.generate_content(prompt)
            response_text = response.text.strip()
            
            # Clean up response - extract JSON from potential markdown
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0].strip()
            elif "```" in response_text:
                response_text = response_text.split("```")[1].split("```")[0].strip()
            
            # Parse JSON
            items = json.loads(response_text)
            
            if not isinstance(items, list):
                items = [items] if isinstance(items, dict) else []
            
            return items
            
        except json.JSONDecodeError as e:
            print(f"JSON parse error from Gemini: {e}")
            return []
        except Exception as e:
            print(f"Gemini extraction error: {e}")
            raise HTTPException(
                status_code=400,
                detail=f"AI extraction failed: {str(e)}. Please use CSV or Excel format."
            )
    
    @staticmethod
    def _validate_items(items: List[dict], workspace_id: str) -> dict:
        """Validate catalogue items"""
        accepted = []
        needs_review = []
        rejected = []
        
        seen_names = set()
        seen_skus = set()
        
        for item in items:
            # Check required fields
            name = str(item.get('name', '')).strip()
            
            if not name:
                rejected.append({"reason": "Missing product/service name", "data": item})
                continue
            
            # Parse price
            price_str = str(item.get('unitPrice', '')).strip()
            price = None
            if price_str:
                try:
                    # Remove currency symbols
                    price_clean = re.sub(r'[£$€,]', '', price_str)
                    price = float(price_clean)
                    if price < 0:
                        needs_review.append({
                            "name": name,
                            "issue": "Negative price",
                            "data": item
                        })
                        continue
                except ValueError:
                    needs_review.append({
                        "name": name,
                        "issue": f"Invalid price format: {price_str}",
                        "data": item
                    })
                    continue
            else:
                needs_review.append({
                    "name": name,
                    "issue": "Missing price",
                    "data": item
                })
                continue
            
            # Check for duplicates
            if name.lower() in seen_names:
                needs_review.append({
                    "name": name,
                    "issue": "Duplicate name - consider merging",
                    "data": item
                })
                continue
            seen_names.add(name.lower())
            
            sku = str(item.get('sku', '')).strip() if item.get('sku') else None
            if sku and sku.lower() in seen_skus:
                needs_review.append({
                    "name": name,
                    "issue": "Duplicate SKU",
                    "data": item
                })
                continue
            if sku:
                seen_skus.add(sku.lower())
            
            # Parse tax rate
            tax_rate = None
            tax_str = str(item.get('taxRate', '')).strip() if item.get('taxRate') else None
            if tax_str:
                try:
                    tax_clean = re.sub(r'[%]', '', tax_str)
                    tax_rate = float(tax_clean)
                except ValueError:
                    pass  # Ignore invalid tax rate
            
            # Build validated item
            validated = {
                "name": name,
                "description": str(item.get('description', '')).strip() or None,
                "unitPrice": price,
                "currency": str(item.get('currency', 'GBP')).strip().upper() or 'GBP',
                "taxRate": tax_rate,
                "sku": sku,
                "category": str(item.get('category', '')).strip() or None
            }
            
            accepted.append(validated)
        
        return {
            "accepted": accepted,
            "needsReview": needs_review,
            "rejected": rejected,
            "summary": {
                "total": len(items),
                "accepted": len(accepted),
                "needsReview": len(needs_review),
                "rejected": len(rejected)
            }
        }
    
    @staticmethod
    async def bulk_add_items(workspace_id: str, user_id: str, items: List[dict]) -> dict:
        """Add multiple validated items to catalogue"""
        db = get_db()
        now = datetime.now(timezone.utc).isoformat()
        
        added_count = 0
        for item_data in items:
            item_id = str(uuid.uuid4())
            
            item = {
                "id": item_id,
                "workspace_id": workspace_id,
                "name": item_data.get("name"),
                "description": item_data.get("description"),
                "unitPrice": item_data.get("unitPrice"),
                "currency": item_data.get("currency", "GBP"),
                "taxRate": item_data.get("taxRate"),
                "sku": item_data.get("sku"),
                "category": item_data.get("category"),
                "created_by": user_id,
                "createdAt": now,
                "updatedAt": now
            }
            
            await db.catalogue_items.insert_one(item)
            added_count += 1
        
        # Log bulk add event
        await db.intelligence_events.insert_one({
            "id": str(uuid.uuid4()),
            "workspace_id": workspace_id,
            "type": "catalogue_bulk_added",
            "data": {"count": added_count},
            "occurredAt": now
        })
        
        return {"success": True, "addedCount": added_count}
