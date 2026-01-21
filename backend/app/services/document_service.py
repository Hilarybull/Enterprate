"""Business Documents Service - Wizard-based Document Generator"""
import uuid
import os
import io
from datetime import datetime, timezone
from typing import Optional, Dict, Any
from fastapi import HTTPException
from pydantic import BaseModel
from app.core.database import get_db

# Try to import LLM for document generation
try:
    from emergentintegrations.llm.chat import LlmChat, UserMessage
    LLM_KEY = os.environ.get("EMERGENT_LLM_KEY")
    LLM_AVAILABLE = bool(LLM_KEY)
except ImportError:
    LLM_AVAILABLE = False
    LLM_KEY = None

# Try to import Gemini as fallback
try:
    import google.generativeai as genai
    GEMINI_KEY = os.environ.get("GEMINI_API_KEY")
    GEMINI_AVAILABLE = bool(GEMINI_KEY)
except ImportError:
    GEMINI_AVAILABLE = False
    GEMINI_KEY = None

# PDF generation with ReportLab (pure Python)
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# DOCX generation
try:
    from docx import Document as DocxDoc
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocumentGenerateRequest(BaseModel):
    documentType: str
    inputs: Dict[str, Any]


class DocumentSaveRequest(BaseModel):
    documentType: str
    title: str
    content: str
    inputs: Dict[str, Any]


class DocumentExportRequest(BaseModel):
    content: str
    title: str
    format: str = "pdf"


class DocumentService:
    """Service for generating business documents via wizard"""
    
    # Document templates and prompts
    DOCUMENT_TEMPLATES = {
        "terms_conditions": {
            "name": "Terms & Conditions",
            "prompt_template": """Generate professional Terms & Conditions for a business with the following details:

Company Name: {companyName}
Business Type: {businessType}
Jurisdiction: {jurisdiction}
Refund Policy: {refundPolicy}
Liability Limit: {liabilityLimit}

Create comprehensive, legally-sound terms covering:
1. Introduction and definitions
2. Use of services/products
3. User obligations
4. Payment terms
5. Refund/cancellation policy (based on provided info)
6. Intellectual property
7. Limitation of liability
8. Termination
9. Governing law ({jurisdiction})
10. Contact information

Format the document professionally with clear headings and numbered sections.
Make it appropriate for {businessType} businesses.
Include today's date as the effective date."""
        },
        "privacy_policy": {
            "name": "Privacy Policy",
            "prompt_template": """Generate a GDPR-compliant Privacy Policy for:

Company Name: {companyName}
Company Address: {companyAddress}
Data Protection Contact: {contactEmail}
Types of Data Collected: {dataCollected}
Third Party Services: {thirdParties}
Data Retention Period: {dataRetention}

The privacy policy must include:
1. Introduction and controller information
2. Data we collect and why
3. Legal basis for processing (GDPR Article 6)
4. How we use your data
5. Data sharing and third parties
6. International transfers
7. Data retention periods
8. Your rights (access, rectification, erasure, portability, objection)
9. Cookies policy
10. Security measures
11. Children's privacy
12. Changes to this policy
13. Contact details for data protection queries

Make it comprehensive but readable. Include specific mention of {thirdParties} if provided."""
        },
        "service_agreement": {
            "name": "Service Agreement",
            "prompt_template": """Generate a professional Service Agreement/Contract between:

Service Provider: {providerName}
Client: {clientName}
Service Description: {serviceDescription}
Key Deliverables: {deliverables}
Project Timeline: {timeline}
Payment Terms: {paymentTerms}
Contract Value: {totalValue}

Include these sections:
1. Parties to the agreement
2. Scope of services
3. Deliverables and milestones
4. Timeline and deadlines
5. Fees and payment schedule
6. Client responsibilities
7. Intellectual property rights
8. Confidentiality
9. Warranties and representations
10. Limitation of liability
11. Termination clauses
12. Force majeure
13. Dispute resolution
14. Signatures block

Make it balanced and fair to both parties. Professional tone."""
        },
        "proposal": {
            "name": "Business Proposal",
            "prompt_template": """Generate a compelling business proposal:

From: {companyName}
To: {clientName}
Project: {projectTitle}

Problem/Opportunity: {problemStatement}
Proposed Solution: {proposedSolution}
Scope of Work: {scope}
Pricing: {pricing}
Timeline: {timeline}

Structure the proposal with:
1. Executive Summary
2. Understanding of the Client's Needs
3. Our Proposed Solution
4. Detailed Scope of Work
5. Methodology/Approach
6. Timeline and Milestones
7. Investment/Pricing
8. Why Choose Us
9. Next Steps
10. Terms and Conditions

Make it persuasive and professional. Focus on value to the client."""
        },
        "nda": {
            "name": "Non-Disclosure Agreement",
            "prompt_template": """Generate a Non-Disclosure Agreement (NDA) between:

Disclosing Party: {disclosingParty}
Receiving Party: {receivingParty}
Purpose: {purpose}
Duration: {duration}
Governing Law: {jurisdiction}

Include:
1. Definition of Confidential Information
2. Obligations of Receiving Party
3. Exclusions from Confidential Information
4. Term of Agreement ({duration})
5. Return of Materials
6. No License Granted
7. No Warranty
8. Remedies
9. Governing Law ({jurisdiction})
10. Entire Agreement
11. Signatures

Make it comprehensive but clear. Balanced obligations."""
        }
    }
    
    @staticmethod
    async def generate_document(workspace_id: str, user_id: str, data: DocumentGenerateRequest) -> dict:
        """Generate a document using AI"""
        db = get_db()
        
        template = DocumentService.DOCUMENT_TEMPLATES.get(data.documentType)
        if not template:
            raise HTTPException(status_code=400, detail="Invalid document type")
        
        # Build prompt from template
        try:
            prompt = template["prompt_template"].format(**{
                k: v or "Not specified" for k, v in data.inputs.items()
            })
        except KeyError as e:
            raise HTTPException(status_code=400, detail=f"Missing required field: {e}")
        
        # Generate document using AI
        if LLM_AVAILABLE:
            try:
                chat = LlmChat(
                    api_key=LLM_KEY,
                    model="anthropic/claude-sonnet-4-20250514"
                )
                
                response = await chat.send_async(
                    UserMessage(content=prompt)
                )
                
                content = response.content
            except Exception:
                # Fallback to basic template
                content = DocumentService._generate_basic_document(data.documentType, data.inputs)
        else:
            content = DocumentService._generate_basic_document(data.documentType, data.inputs)
        
        # Generate title
        title = DocumentService._generate_title(data.documentType, data.inputs)
        
        # Log generation event
        await db.intelligence_events.insert_one({
            "id": str(uuid.uuid4()),
            "workspace_id": workspace_id,
            "type": "document_generated",
            "data": {
                "documentType": data.documentType,
                "title": title
            },
            "occurredAt": datetime.now(timezone.utc).isoformat()
        })
        
        return {
            "title": title,
            "content": content,
            "documentType": data.documentType,
            "generatedAt": datetime.now(timezone.utc).isoformat()
        }
    
    @staticmethod
    def _generate_title(doc_type: str, inputs: dict) -> str:
        """Generate document title"""
        template = DocumentService.DOCUMENT_TEMPLATES.get(doc_type, {})
        name = template.get("name", "Document")
        
        if doc_type == "terms_conditions":
            return f"Terms & Conditions - {inputs.get('companyName', 'Company')}"
        elif doc_type == "privacy_policy":
            return f"Privacy Policy - {inputs.get('companyName', 'Company')}"
        elif doc_type == "service_agreement":
            return f"Service Agreement - {inputs.get('providerName', 'Provider')} & {inputs.get('clientName', 'Client')}"
        elif doc_type == "proposal":
            return f"Proposal - {inputs.get('projectTitle', 'Project')}"
        elif doc_type == "nda":
            return f"NDA - {inputs.get('disclosingParty', 'Party A')} & {inputs.get('receivingParty', 'Party B')}"
        
        return name
    
    @staticmethod
    def _generate_basic_document(doc_type: str, inputs: dict) -> str:
        """Generate a basic document without AI"""
        now = datetime.now().strftime("%d %B %Y")
        
        if doc_type == "terms_conditions":
            return f"""TERMS AND CONDITIONS

Effective Date: {now}

1. INTRODUCTION

These Terms and Conditions govern your use of services provided by {inputs.get('companyName', '[Company Name]')}.

2. SERVICES

We provide {inputs.get('businessType', 'business')} services as described on our website.

3. PAYMENT TERMS

Payment is due as specified in your order or invoice.

4. REFUND POLICY

{inputs.get('refundPolicy', 'Refunds may be provided at our discretion.')}

5. LIMITATION OF LIABILITY

Our liability is limited to {inputs.get('liabilityLimit', 'the amount paid for services')}.

6. GOVERNING LAW

These terms are governed by the laws of {inputs.get('jurisdiction', 'England & Wales')}.

7. CONTACT

For questions about these terms, please contact us.

Last updated: {now}
"""
        
        elif doc_type == "privacy_policy":
            return f"""PRIVACY POLICY

Effective Date: {now}

1. WHO WE ARE

{inputs.get('companyName', '[Company Name]')}
{inputs.get('companyAddress', '[Address]')}

2. DATA WE COLLECT

{inputs.get('dataCollected', 'We collect personal information necessary to provide our services.')}

3. HOW WE USE YOUR DATA

We use your data to provide and improve our services.

4. THIRD PARTIES

{inputs.get('thirdParties', 'We may share data with service providers who assist us.')}

5. DATA RETENTION

{inputs.get('dataRetention', 'We retain data as long as necessary for business purposes.')}

6. YOUR RIGHTS

You have the right to access, correct, or delete your personal data.

7. CONTACT US

Data Protection Contact: {inputs.get('contactEmail', '[email]')}

Last updated: {now}
"""
        
        # Add more templates as needed
        return f"Document generated on {now}\n\nPlease customize this document for your needs."
    
    @staticmethod
    async def save_document(workspace_id: str, user_id: str, data: DocumentSaveRequest) -> dict:
        """Save a generated document"""
        db = get_db()
        
        doc_id = str(uuid.uuid4())
        now = datetime.now(timezone.utc).isoformat()
        
        document = {
            "id": doc_id,
            "workspace_id": workspace_id,
            "documentType": data.documentType,
            "title": data.title,
            "content": data.content,
            "inputs": data.inputs,
            "version": 1,
            "created_by": user_id,
            "createdAt": now,
            "updatedAt": now
        }
        
        await db.documents.insert_one(document)
        
        # Log save event
        await db.intelligence_events.insert_one({
            "id": str(uuid.uuid4()),
            "workspace_id": workspace_id,
            "type": "document_saved",
            "data": {
                "documentId": doc_id,
                "title": data.title,
                "documentType": data.documentType
            },
            "occurredAt": now
        })
        
        return {"id": doc_id, "success": True}
    
    @staticmethod
    async def get_documents(workspace_id: str, doc_type: Optional[str] = None) -> list:
        """Get all documents for workspace"""
        db = get_db()
        
        query = {"workspace_id": workspace_id}
        if doc_type:
            query["documentType"] = doc_type
        
        docs = await db.documents.find(query).sort("createdAt", -1).to_list(length=100)
        return [{k: v for k, v in doc.items() if k != '_id'} for doc in docs]
    
    @staticmethod
    async def get_document(doc_id: str, workspace_id: str) -> dict:
        """Get a single document"""
        db = get_db()
        
        doc = await db.documents.find_one({
            "id": doc_id,
            "workspace_id": workspace_id
        })
        
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        
        return {k: v for k, v in doc.items() if k != '_id'}
    
    @staticmethod
    async def export_document(data: DocumentExportRequest) -> bytes:
        """Export document to PDF or DOCX"""
        
        if data.format == "pdf":
            return DocumentService._generate_pdf(data.title, data.content)
        elif data.format == "docx":
            return DocumentService._generate_docx(data.title, data.content)
        else:
            # Plain text fallback
            content = f"{data.title}\n\n{'='*50}\n\n{data.content}"
            return content.encode('utf-8')
    
    @staticmethod
    def _generate_pdf(title: str, content: str) -> bytes:
        """Generate PDF from document content using ReportLab"""
        
        if REPORTLAB_AVAILABLE:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(
                buffer, 
                pagesize=A4, 
                topMargin=25*mm, 
                bottomMargin=25*mm,
                leftMargin=20*mm,
                rightMargin=20*mm
            )
            
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'DocTitle',
                parent=styles['Title'],
                fontSize=24,
                textColor=colors.HexColor('#4F46E5'),
                spaceAfter=20
            )
            
            heading_style = ParagraphStyle(
                'DocHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.HexColor('#333333'),
                spaceBefore=15,
                spaceAfter=8
            )
            
            body_style = ParagraphStyle(
                'DocBody',
                parent=styles['Normal'],
                fontSize=11,
                leading=16,
                spaceAfter=8
            )
            
            date_style = ParagraphStyle(
                'DocDate',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.grey,
                alignment=1  # Center
            )
            
            elements = []
            
            # Title
            elements.append(Paragraph(title, title_style))
            elements.append(Paragraph(f"Generated on {datetime.now().strftime('%d %B %Y')}", date_style))
            elements.append(Spacer(1, 20))
            
            # Content - split by paragraphs and detect headings
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                
                # Check if it's a heading (numbered section or all caps)
                lines = para.split('\n')
                first_line = lines[0].strip()
                
                if first_line and (
                    (first_line[0].isdigit() and '. ' in first_line[:5] and len(first_line) < 80) or
                    (first_line.isupper() and len(first_line) < 60)
                ):
                    # It's a heading
                    elements.append(Paragraph(first_line, heading_style))
                    # Add remaining lines as body
                    if len(lines) > 1:
                        remaining = '\n'.join(lines[1:]).strip()
                        if remaining:
                            elements.append(Paragraph(remaining.replace('\n', '<br/>'), body_style))
                else:
                    # Regular paragraph
                    elements.append(Paragraph(para.replace('\n', '<br/>'), body_style))
            
            doc.build(elements)
            return buffer.getvalue()
        
        else:
            # Plain text with header as fallback
            content_str = f"{title}\n{'='*50}\nGenerated on {datetime.now().strftime('%d %B %Y')}\n{'='*50}\n\n{content}"
            return content_str.encode('utf-8')
    
    @staticmethod
    def _generate_docx(title: str, content: str) -> bytes:
        """Generate DOCX from document content"""
        
        if not DOCX_AVAILABLE:
            raise HTTPException(status_code=500, detail="DOCX generation not available")
        
        doc = DocxDoc()
        
        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_para = doc.add_paragraph(f"Generated on {datetime.now().strftime('%d %B %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Empty line
        
        # Add content
        for para in content.split('\n\n'):
            if para.strip():
                # Check if it's a heading (starts with number and period or all caps)
                stripped = para.strip()
                if stripped and (stripped[0].isdigit() and '. ' in stripped[:5]):
                    # It's likely a section heading
                    doc.add_heading(stripped, level=1)
                elif stripped.isupper() and len(stripped) < 100:
                    # All caps - likely a heading
                    doc.add_heading(stripped, level=1)
                else:
                    doc.add_paragraph(stripped)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    
    @staticmethod
    def _format_content_for_html(content: str) -> str:
        """Format plain text content for HTML display"""
        import html
        
        # Escape HTML entities
        content = html.escape(content)
        
        # Convert line breaks
        content = content.replace('\n\n', '</p><p>')
        content = content.replace('\n', '<br>')
        
        # Wrap in paragraph
        content = f"<p>{content}</p>"
        
        # Format headings (lines that look like section titles)
        import re
        # Numbers followed by period and text
        content = re.sub(r'<p>(\d+)\.\s+([A-Z][A-Z\s]+)(?:</p>|<br>)', r'<h2>\1. \2</h2>', content)
        
        return content
